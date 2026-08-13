import os
import logging
from typing import Dict, Any

logger = logging.getLogger("placex.docx_extractor")

class DOCXExtractor:
    """
    Production-Grade DOCX Text Extractor for PlaceX using python-docx.
    """

    @classmethod
    def extract_text_and_info(cls, file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            return {
                "text": "",
                "status": "EXTRACTION_FAILED",
                "character_count": 0,
                "error": f"DOCX file not found at {file_path}"
            }

        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    full_text.append(para.text.strip())
            
            # Extract text from tables as well
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            text = "\n".join(full_text).strip()
            char_count = len(text)
            status = "SUCCESS" if char_count > 0 else "EXTRACTION_FAILED"

            return {
                "text": text,
                "status": status,
                "character_count": char_count,
                "error": None if status == "SUCCESS" else "DOCX contains no readable text."
            }
        except Exception as e:
            logger.error(f"DOCX extraction error for {file_path}: {e}")
            return {
                "text": "",
                "status": "EXTRACTION_FAILED",
                "character_count": 0,
                "error": str(e)
            }
